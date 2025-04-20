# logic/a_resume/uploadResume.py - Handles resume upload and parsing
import os
import tempfile
import json
import logging
from flask import Blueprint, request, redirect, url_for, flash
from werkzeug.utils import secure_filename
from app_logic.a_resume.resumeHistory import resume_storage
from app_logic.b_jobs.jobMatch import generate_dual_embeddings
logger = logging.getLogger(__name__)
upload_resume_bp = Blueprint("upload_resume", __name__)
# === Configuration ===
ADZUNA_DATA_DIR = os.path.join(os.path.dirname(__file__), '../../static/job_data/adzuna')
ADZUNA_INDEX_FILE = os.path.join(ADZUNA_DATA_DIR, 'index.json')
RESUME_INDEX_FILE = os.path.join(os.path.dirname(__file__), '../../static/resumes/index.json')
ALLOWED_EXTENSIONS = {"docx", "txt"}
TEMP_FOLDER = tempfile.gettempdir()

# === Resume Parsing ===
class FileParsingError(Exception):
    pass
# """Extracts text from DOCX file including paragraphs, tables, headers, footers"""
def parse_docx(file_path):
    try:
        from docx import Document
    except ImportError as e:
        logger.error(f"DOCX parser unavailable: {str(e)}")
        raise FileParsingError("DOCX parsing is not available. Please upload a TXT file instead.")
    try:
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = " ".join(p.text for p in cell.paragraphs if p.text.strip())
                    if cell_text:
                        row_text.append(cell_text.strip())
                if row_text:
                    full_text.append(" | ".join(row_text))
        for section in doc.sections:
            for paragraph in section.header.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)
            for paragraph in section.footer.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)
        text = "\n".join(full_text).strip()
        if not text:
            raise FileParsingError("The DOCX file appears to be empty")
        logger.info("Successfully parsed DOCX file: %s", file_path)
        return text
    except Exception as e:
        logger.error(f"Error parsing DOCX: {str(e)}")
        raise FileParsingError(f"Failed to parse DOCX: {str(e)}")
# """Extracts text from a plain text file"""
def parse_txt(file_path):
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read().strip()
        if not text:
            raise FileParsingError("The text file appears to be empty")
        return text
    except Exception as e:
        logger.error(f"Error reading TXT: {str(e)}")
        raise FileParsingError(f"Failed to read text file: {str(e)}")
# """Dispatch parsing based on file extension"""
def parse_resume(file_path):
    if not os.path.exists(file_path):
        raise FileParsingError(f"File not found: {file_path}")
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".docx":
        return parse_docx(file_path)
    elif ext == ".txt":
        return parse_txt(file_path)
    else:
        raise FileParsingError(f"Unsupported file type: {ext}")
# === Internal Utility ===
# """Return True if the filename is a supported file type"""
def allowed_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS
# """Save the resume index with embedded arrays converted to lists"""
def export_resume_index_with_embeddings(resume_storage_instance) -> None:
    try:
        index_copy = {
            "resumes": {},
            "count": resume_storage_instance._index["count"],
            "last_added": resume_storage_instance._index["last_added"]
        }
        for resume_id, resume_data in resume_storage_instance._index["resumes"].items():
            resume_copy = resume_data.copy()
            if "embedding" in resume_copy and hasattr(resume_copy["embedding"], "tolist"):
                resume_copy["embedding"] = resume_copy["embedding"].tolist()
            index_copy["resumes"][resume_id] = resume_copy
        with open(RESUME_INDEX_FILE, 'w', encoding='utf-8') as f:
            json.dump(index_copy, f, indent=2)
        logger.info("Exported resume index with embeddings serialized")
    except Exception as e:
        logger.error(f"Error saving resume index: {str(e)}")


# === Upload Resume Route ===
# """Route for handling resume file upload and processing"""
@upload_resume_bp.route("/upload_resume", methods=["POST"])
def upload_resume():
    if "resume" not in request.files:
        flash("No file part", "danger")
        return redirect(url_for("index"))
    file = request.files["resume"]
    if file.filename == "":
        flash("No file selected", "danger")
        return redirect(url_for("index"))
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        filepath = os.path.join(TEMP_FOLDER, filename)
        file.save(filepath)
        try:
            resume_text = parse_resume(filepath)
        except FileParsingError as e:
            flash(f"Resume parsing error: {str(e)}", "danger")
            return redirect(url_for("index"))
        except Exception as e:
            logger.exception("Unexpected resume parsing error")
            flash(f"Unexpected error: {str(e)}", "danger")
            return redirect(url_for("index"))
        metadata = {
            "filters": {
                "remote": request.form.get("remote", "") == "on",
                "location": request.form.get("location", ""),
                "keywords": request.form.get("keywords", "")
            }
        }
        embeddings = generate_dual_embeddings(resume_text)
        metadata["embedding_narrative"] = embeddings["narrative"].tolist()
        metadata["embedding_skills"] = embeddings["skills"].tolist()
        resume_id = resume_storage.store_resume(
            temp_filepath=filepath,
            filename=filename,
            content=resume_text,
            metadata=metadata
        )
        flash(f'Resume "{filename}" successfully uploaded and stored', "success")
        return redirect(url_for("index"))
    flash("Invalid file type", "danger")
    return redirect(url_for("index"))
