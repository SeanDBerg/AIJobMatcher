# client/resume/uploadResume.py - Handles resume upload and parsing
import os
import tempfile
import logging
from flask import Blueprint, request, redirect, url_for, flash
from werkzeug.utils import secure_filename
from matching_engine import generate_dual_embeddings
from resume_storage import resume_storage
from job_manager import JobManager
logger = logging.getLogger(__name__)
upload_resume_bp = Blueprint("upload_resume", __name__)

ALLOWED_EXTENSIONS = {"docx", "txt"}
TEMP_FOLDER = tempfile.gettempdir()
job_manager = JobManager()

# === Resume Parsing ===

class FileParsingError(Exception):
    """Custom exception for file parsing errors"""
    pass

def parse_docx(file_path):
    """Extracts text from DOCX file including paragraphs, tables, headers, footers"""
    try:
        from docx import Document
    except ImportError as e:
        logger.error(f"DOCX parser unavailable: {str(e)}")
        raise FileParsingError("DOCX parsing is not available. Please upload a TXT file instead.")

    try:
        doc = Document(file_path)
        full_text = []

        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = " ".join(p.text for p in cell.paragraphs if p.text.strip())
                    if cell_text:
                        row_text.append(cell_text.strip())
                if row_text:
                    full_text.append(" | ".join(row_text))

        for section in doc.sections:
            for paragraph in section.header.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)
            for paragraph in section.footer.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)

        text = "\n".join(full_text).strip()
        if not text:
            raise FileParsingError("The DOCX file appears to be empty")

        logger.info("Successfully parsed DOCX file: %s", file_path)
        return text

    except Exception as e:
        logger.error(f"Error parsing DOCX: {str(e)}")
        raise FileParsingError(f"Failed to parse DOCX: {str(e)}")

def parse_txt(file_path):
    """Extracts text from a plain text file"""
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read().strip()
        if not text:
            raise FileParsingError("The text file appears to be empty")
        return text
    except Exception as e:
        logger.error(f"Error reading TXT: {str(e)}")
        raise FileParsingError(f"Failed to read text file: {str(e)}")

def parse_resume(file_path):
    """Dispatch parsing based on file extension"""
    if not os.path.exists(file_path):
        raise FileParsingError(f"File not found: {file_path}")

    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".docx":
        return parse_docx(file_path)
    elif ext == ".txt":
        return parse_txt(file_path)
    else:
        raise FileParsingError(f"Unsupported file type: {ext}")

# === Internal Utility ===
def allowed_file(filename: str) -> bool:
    """Return True if the filename is a supported file type"""
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

# === Upload Resume Route ===

@upload_resume_bp.route("/upload_resume", methods=["POST"])
def upload_resume():
    """Route for handling resume file upload and processing"""

    if "resume" not in request.files:
        flash("No file part", "danger")
        return redirect(url_for("index"))

    file = request.files["resume"]

    if file.filename == "":
        flash("No file selected", "danger")
        return redirect(url_for("index"))

    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        filepath = os.path.join(TEMP_FOLDER, filename)
        file.save(filepath)

        try:
            resume_text = parse_resume(filepath)
        except FileParsingError as e:
            flash(f"Resume parsing error: {str(e)}", "danger")
            return redirect(url_for("index"))
        except Exception as e:
            logger.exception("Unexpected resume parsing error")
            flash(f"Unexpected error: {str(e)}", "danger")
            return redirect(url_for("index"))

        try:
            embeddings = generate_dual_embeddings(resume_text)
        except Exception as e:
            logger.exception("Error generating embeddings")
            flash(f"Error analyzing resume content: {str(e)}", "danger")
            return redirect(url_for("index"))

        filters = {
            "remote": request.form.get("remote", "") == "on",
            "location": request.form.get("location", ""),
            "keywords": request.form.get("keywords", "")
        }

        metadata = {
            "embedding_narrative": embeddings["narrative"].tolist(),
            "embedding_skills": embeddings["skills"].tolist(),
            "filters": filters
        }

        resume_id = resume_storage.store_resume(
            temp_filepath=filepath,
            filename=filename,
            content=resume_text,
            metadata=metadata
        )

        flash(f'Resume "{filename}" successfully uploaded and stored', "success")

        if request.form.get("find_matches") == "on":
            try:
                jobs = job_manager.get_recent_jobs(days=30)
                if not jobs:
                    flash("No job data available to match against", "warning")
                    return redirect(url_for("index", resume_id=resume_id))
            except Exception as e:
                logger.exception("Error retrieving job data")
                flash(f"Error retrieving job data: {str(e)}", "danger")
                return redirect(url_for("index", resume_id=resume_id))

        return redirect(url_for("index", resume_id=resume_id))

    flash("Invalid file type", "danger")
    return redirect(url_for("index"))
