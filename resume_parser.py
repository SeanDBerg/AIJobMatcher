import logging
import os
from io import StringIO

logger = logging.getLogger(__name__)

class FileParsingError(Exception):
    """Custom exception for file parsing errors"""
    pass
# Parse a PDF file and extract text content
def parse_pdf(file_path):
    logger.debug(f"Attempting to parse PDF: {file_path}")
    try:
        # Try to import PDF parsing libraries
        try:
            from pdfminer.converter import TextConverter
            from pdfminer.layout import LAParams
            from pdfminer.pdfdocument import PDFDocument
            from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
            from pdfminer.pdfpage import PDFPage
            from pdfminer.pdfparser import PDFParser
        except ImportError as e:
            logger.error(f"PDF parsing libraries not available: {str(e)}")
            raise FileParsingError(
                "PDF parsing libraries are not available. The system cannot process PDF files at this time. "
                "Please upload a text (.txt) file instead."
            )
        output_string = StringIO()
        try:
            with open(file_path, 'rb') as file:
                parser = PDFParser(file)
                doc = PDFDocument(parser)
                rsrcmgr = PDFResourceManager()
                device = TextConverter(rsrcmgr, output_string, laparams=LAParams())
                interpreter = PDFPageInterpreter(rsrcmgr, device)
                for page in PDFPage.create_pages(doc):
                    interpreter.process_page(page)
            text = output_string.getvalue()
            output_string.close()
            if not text.strip():
                raise FileParsingError("The PDF file appears to be empty or contains no extractable text")
            return text.strip()
        except Exception as e:
            logger.error(f"Error parsing PDF: {str(e)}")
            raise FileParsingError(f"Failed to extract text from the PDF file: {str(e)}")
    except Exception as e:
        if isinstance(e, FileParsingError):
            raise
        logger.error(f"Unexpected error in PDF parsing: {str(e)}")
        raise FileParsingError(f"An unexpected error occurred while parsing the PDF: {str(e)}")
# Parse a DOCX file and extract text content
def parse_docx(file_path):
    logger.debug(f"Attempting to parse DOCX: {file_path}")
    try:
        # Try to import DOCX parsing libraries
        try:
            from docx import Document
        except ImportError as e:
            logger.error(f"DOCX parsing libraries not available: {str(e)}")
            raise FileParsingError(
                "DOCX parsing libraries are not available. The system cannot process DOCX files at this time. "
                "Please upload a text (.txt) file instead."
            )
        try:
            doc = Document(file_path)
            full_text = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():  # Only add non-empty paragraphs
                    full_text.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        # Get text from paragraphs in cells
                        cell_text = ""
                        for paragraph in cell.paragraphs:
                            if paragraph.text.strip():
                                cell_text += paragraph.text + " "
                        if cell_text.strip():
                            row_text.append(cell_text.strip())
                    if row_text:
                        full_text.append(" | ".join(row_text))
            
            # Extract text from headers and footers
            for section in doc.sections:
                # Header
                for paragraph in section.header.paragraphs:
                    if paragraph.text.strip():
                        full_text.append(paragraph.text)
                
                # Footer
                for paragraph in section.footer.paragraphs:
                    if paragraph.text.strip():
                        full_text.append(paragraph.text)
            
            text = '\n'.join(full_text).strip()
            if not text:
                raise FileParsingError("The DOCX file appears to be empty or contains no extractable text")
            
            logger.info(f"Successfully parsed DOCX file: {file_path}")
            return text
        except Exception as e:
            logger.error(f"Error parsing DOCX: {str(e)}")
            raise FileParsingError(f"Failed to extract text from the DOCX file: {str(e)}")
    except Exception as e:
        if isinstance(e, FileParsingError):
            raise
        logger.error(f"Unexpected error in DOCX parsing: {str(e)}")
        raise FileParsingError(f"An unexpected error occurred while parsing the DOCX: {str(e)}")
# Parse a TXT file and extract text content
def parse_txt(file_path):
    logger.debug(f"Parsing TXT: {file_path}")
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            text = file.read()
        text = text.strip()
        if not text:
            raise FileParsingError("The text file appears to be empty")
        return text
    except FileParsingError:
        raise
    except Exception as e:
        logger.error(f"Error reading text file: {str(e)}")
        raise FileParsingError(f"Failed to read the text file: {str(e)}")
# Parse a resume file and extract text content based on file extension
def parse_resume(file_path):
    logger.debug(f"Parsing resume: {file_path}")
    if not os.path.exists(file_path):
        raise FileParsingError(f"The file does not exist: {file_path}")
    file_extension = os.path.splitext(file_path)[1].lower()
    if file_extension == '.pdf':
        resume_text = parse_pdf(file_path)
    elif file_extension == '.docx':
        resume_text = parse_docx(file_path)
    elif file_extension == '.txt':
        resume_text = parse_txt(file_path)
    else:
        raise ValueError(f"Unsupported file extension: {file_extension}. Please upload a PDF, DOCX, or TXT file.")
    logger.info(f"Parsed resume text: {resume_text}")
    return resume_text
